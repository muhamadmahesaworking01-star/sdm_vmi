from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

out = 'docs/Laporan-Poin-Penting-SI-Villa-Merah.docx'
doc = Document(); sec = doc.sections[0]
sec.top_margin = Inches(.75); sec.bottom_margin = Inches(.75); sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
doc.styles['Normal'].font.name='Aptos'; doc.styles['Normal'].font.size=Pt(10)
for name,size,color in [('Title',26,'17365D'),('Heading 1',16,'1F4E79'),('Heading 2',12,'2E75B6')]:
 s=doc.styles[name]; s.font.name='Aptos Display'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
def bullet(text): doc.add_paragraph(text, style='List Bullet')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('LAPORAN POIN PENTING\nSISTEM INFORMASI VILLA MERAH'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(23,54,93)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Ringkasan kebutuhan pengembangan, arsitektur, keamanan, dan operasional').italic=True
doc.add_paragraph('Dokumen ini merangkum daftar isi dan poin penting dokumen pengembangan Sistem Informasi Villa Merah serta implementasi aplikasi.', style='Intense Quote')
doc.add_heading('1. Gambaran Umum Sistem', level=1); doc.add_paragraph('Sistem Informasi Villa Merah adalah aplikasi pengelolaan sumber daya manusia berbasis Laravel dengan role-based access control. Sistem memusatkan data akun, pegawai, pengajar, dokumen, kompetensi, pengumuman, audit aktivitas, dan pencadangan data.')
doc.add_heading('2. Peran dan Hak Akses', level=1)
t=doc.add_table(rows=1, cols=3); t.style='Light Shading Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for c,v in zip(t.rows[0].cells,['Peran','Fokus','Hak utama']): c.text=v
for row in [('Super Admin','Administrasi sistem','Kelola akun, pegawai, pengajar, pengumuman, backup, log, biodata sendiri'),('Karyawan','Self-service kantor','Profil dan dokumen sendiri'),('Pengajar','Self-service akademik','Profil, kampus, kompetensi, dokumen sendiri'),('Karyawan & Pengajar','Self-service hybrid','Profil struktural dan akademik'),('Direksi','Monitoring','Ringkasan organisasi dan direktori read-only')]:
 for c,v in zip(t.add_row().cells,row): c.text=v
doc.add_heading('3. Modul Aplikasi', level=1)
for x in ['Autentikasi dan manajemen akun','Dashboard per role dan kelengkapan profil','Biodata self-service seluruh pegawai dan Super Admin','Manajemen dokumen dan berkas','Master spesialisasi pengajar','Pengumuman internal','Monitoring Direksi','Impor/ekspor Excel','Log aktivitas user','Backup ZIP JSON']: bullet(x)
doc.add_heading('4. Data, Relasi, dan Kardinalitas', level=1); doc.add_paragraph('Entitas inti: users, employees, employee_documents, specializations, announcements, activity_logs, payroll, salary_components, dan contract_histories. Users memiliki satu role; pegawai memiliki banyak dokumen dan relasi HR berbasis NIP; activity_logs terhubung opsional ke users. Backup mengambil snapshot tabel inti.')
doc.add_heading('5. Alur Proses', level=1)
for x in ['Login → autentikasi → middleware role → dashboard sesuai role.','Pegawai memilih Edit Profil → mengisi data → Simpan Perubahan Profil → data tersimpan dan ditampilkan kembali dari database.','Super Admin mengisi Biodata Saya dengan alur yang sama.','Super Admin menjalankan Backup → snapshot tabel → ZIP private → unduhan terautentikasi.','Request terautentikasi dicatat middleware LogActivity.']: bullet(x)
doc.add_heading('6. Keamanan dan Audit', level=1)
for x in ['Role middleware membatasi route.','Staf hanya mengubah profil sendiri.','Direksi tidak memiliki route CRUD.','Form memakai CSRF dan validasi server-side.','Impor Excel transaksional dan menolak baris invalid.','Backup hanya untuk Super Admin di storage private.','Aktivitas dicatat pada activity_logs.']: bullet(x)
doc.add_heading('7. Backup, Restore, API, dan Teknologi', level=1); doc.add_paragraph('Backup menghasilkan arsip ZIP berisi manifest dan JSON tabel inti. Restore dilakukan setelah verifikasi arsip, backup database aktif, validasi struktur, dan uji integritas relasi. Endpoint internal Laravel menggunakan session authentication, CSRF, dan role middleware. Stack: PHP 8.3, Laravel 13, Blade, Bootstrap/CSS, Eloquent, PhpSpreadsheet, dan ZipArchive.')
doc.add_heading('8. Checklist Laporan', level=1)
for x in ['ERD dan kardinalitas','DFD dan flow proses','Use case dan hak akses','Struktur tabel database','Modul aplikasi','Enterprise architecture','API dan integrasi','Security checklist','Backup dan restore','Fitur aplikasi']: bullet(x)
doc.add_paragraph('Sumber: Daftar Isi dan Poin Penting dalam Dokumen Pengembangan SISTEM INFORMASI VILLA MERAH.pdf serta dokumentasi proyek pada folder docs.', style='Intense Quote')
doc.save(out); print(out)
