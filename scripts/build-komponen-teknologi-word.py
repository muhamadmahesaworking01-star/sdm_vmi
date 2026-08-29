from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT=Path('docs/Daftar-Komponen-Teknologi-SI-SDM.docx'); BLUE='2E74B5'; DARK='1F4D78'; LIGHT='F2F4F7'; INK='0B2545'; MUTED='666666'
def shade(c,fill):
    p=c._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); p.append(x)
def widths(t,ws):
    t.autofit=False; pr=t._tbl.tblPr
    for tag,val in [('w:tblW',sum(ws)),('w:tblInd',120)]:
        x=pr.find(qn(tag)) or OxmlElement(tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa')
        if x.getparent() is None: pr.append(x)
    g=t._tbl.tblGrid
    for x in list(g): g.remove(x)
    for w in ws: x=OxmlElement('w:gridCol'); x.set(qn('w:w'),str(w)); g.append(x)
    for row in t.rows:
        for c,w in zip(row.cells,ws):
            p=c._tc.get_or_add_tcPr(); x=OxmlElement('w:tcW'); x.set(qn('w:w'),str(w)); x.set(qn('w:type'),'dxa'); p.append(x)
            m=OxmlElement('w:tcMar')
            for s,v in [('top',90),('start',120),('bottom',90),('end',120)]: z=OxmlElement('w:'+s); z.set(qn('w:w'),str(v)); z.set(qn('w:type'),'dxa'); m.append(z)
            p.append(m); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(d,heads,rows,ws):
    t=d.add_table(rows=1,cols=len(heads)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT; widths(t,ws)
    for c,h in zip(t.rows[0].cells,heads): shade(c,LIGHT); c.text=h; c.paragraphs[0].runs[0].bold=True
    for row in rows:
        cells=t.add_row().cells
        for c,v in zip(cells,row): c.text=str(v)
    d.add_paragraph()
def label(d,k,v):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(6); r=p.add_run(k); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK); p.add_run(v)

d=Document(); s=d.sections[0]
for a in ['top_margin','bottom_margin','left_margin','right_margin']: setattr(s,a,Inches(1))
s.header_distance=Inches(.492); s.footer_distance=Inches(.492)
n=d.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.1
for nm,sz,col,bef,aft in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
    st=d.styles[nm]; st.font.name='Calibri'; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(col); st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft)
h=s.header.paragraphs[0]; h.text='SI SDM | DAFTAR KOMPONEN TEKNOLOGI'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
f=s.footer.paragraphs[0]; f.text='Dokumentasi teknis proyek | 27 Agustus 2026'; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.runs[0].font.size=Pt(8); f.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
p=d.add_paragraph(); r=p.add_run('DAFTAR KOMPONEN TEKNOLOGI'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string(INK)
p=d.add_paragraph(); r=p.add_run('Sistem Informasi Sumber Daya Manusia (SI SDM)'); r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(BLUE); p.paragraph_format.space_after=Pt(16)
label(d,'Tujuan dokumen: ','Menyajikan komponen teknologi yang membentuk SI SDM, fungsi masing-masing komponen, dan keterkaitannya dalam arsitektur aplikasi.')
label(d,'Ruang lingkup: ','Disusun berdasarkan struktur kode, konfigurasi, dan dokumentasi proyek yang tersedia.')
d.add_heading('1. Ringkasan Komponen Utama',1)
table(d,['Kelompok','Komponen','Peran dalam SI SDM'],[
('Platform','PHP 8.4; Laravel 13.8','Runtime dan framework utama aplikasi web.'),
('Antarmuka','Blade; CSS/Tailwind; Vite','Dashboard, formulir, tabel, dan aset frontend.'),
('Data','Database relasional; Eloquent ORM','Penyimpanan dan akses data domain SDM.'),
('Keamanan','Session; CSRF; role middleware; validasi','Autentikasi, otorisasi, dan perlindungan perubahan data.'),
('Dokumen/laporan','Private storage; ZipArchive; PhpSpreadsheet','Berkas pegawai, backup ZIP, impor/ekspor XLSX.'),
('Kualitas/operasi','PHPUnit; Laravel Pint; deploy.sh','Pengujian, konsistensi kode, dan rilis.')],[1700,2700,4960])
d.add_heading('2. Komponen Teknologi Berdasarkan Lapisan',1)
d.add_heading('2.1 Presentasi / Frontend',2)
table(d,['Komponen','Teknologi / artefak','Fungsi'],[
('Template UI','Laravel Blade','Halaman login, dashboard, profil, kontrak, dokumen, dan admin.'),
('Styling','CSS dan Tailwind CSS','Layout, tipografi, warna, responsivitas, dan komponen visual.'),
('Build asset','Vite dan laravel-vite-plugin','Kompilasi dan pengelolaan aset CSS/JavaScript.'),
('Navigasi role','Blade layouts, sidebar, route grouping','Menu dan dashboard sesuai role pengguna.')],[1800,3000,4560])
d.add_heading('2.2 Aplikasi / Backend',2)
table(d,['Komponen','Teknologi / artefak','Fungsi'],[
('Arsitektur','Laravel MVC','Memisahkan route, middleware, controller, model, dan view.'),
('Routing','routes/web.php','Endpoint internal berbasis session dan CSRF.'),
('Middleware','auth dan role middleware','Akses super_admin, direksi, karyawan, pengajar, dan karyawan_pengajar.'),
('Business logic','Controller dan service','Use case SDM: pegawai, dokumen, kontrak, payroll, pengumuman, backup, laporan.'),
('ORM','Laravel Eloquent','Pemetaan model aplikasi ke tabel dan operasi CRUD.')],[1800,3000,4560])
d.add_heading('2.3 Data dan Penyimpanan',2)
table(d,['Komponen','Objek / teknologi','Fungsi'],[
('Data akun & SDM','users; employees; specializations','Identitas login, profil pegawai, dan master kompetensi.'),
('Data operasional','employee_documents; announcements; activity_logs','Dokumen, siaran internal, dan audit.'),
('Kompensasi','Komponen gaji; payroll bulanan','Penghasilan, bonus, slip, dan status pembayaran.'),
('Kontrak','Riwayat kontrak','Periode, tipe, perpanjangan, dan keterangan kontrak.'),
('File & backup','Private disk; JSON snapshot; ZipArchive','Berkas terlindungi dan backup tabel dalam ZIP.')],[1800,3000,4560])
d.add_heading('3. Modul Fungsional dan Komponen Pendukung',1)
table(d,['Modul','Komponen terkait','Kemampuan'],[
('Autentikasi & akun','Session, password hashing, reset token','Login, logout, reset password, status akun, impersonasi.'),
('Manajemen SDM','Employee model, controller admin, import/export','CRUD pegawai/pengajar, status, template impor, ekspor.'),
('Profil & dokumen','Validasi, private storage','Pembaruan profil serta unggah/unduh dokumen.'),
('Kompetensi','Specializations, kompetensi, portofolio','Master keahlian dan data akademik pengajar.'),
('Kontrak & payroll','Controller kontrak, endpoint JSON internal','Monitoring, kontrak berakhir, riwayat, perpanjangan, ringkasan.'),
('Pengumuman & audit','Announcements; LogActivity; activity_logs','Siaran internal dan pencatatan user, route, metode, IP, waktu.')],[1800,3000,4560])
d.add_heading('4. Keamanan dan Pengendalian Akses',1)
label(d,'Otorisasi berbasis peran: ','Akses dipisahkan melalui middleware role. Direksi bersifat read-only; Super Admin menangani administrasi dan backup.')
label(d,'Perlindungan formulir: ','Request POST, PUT, PATCH, dan DELETE memakai CSRF token serta validasi server-side.')
label(d,'Perlindungan data: ','Backup berada di storage private dan dilayani melalui response terautentikasi. Profil staf ditentukan dari identitas pengguna aktif.')
label(d,'Auditabilitas: ','Aktivitas request dicatat untuk penelusuran perubahan dan pemeriksaan operasional.')
d.add_heading('5. Pengembangan, Pengujian, dan Deployment',1)
table(d,['Area','Komponen','Keterangan'],[
('Backend','Composer; composer.lock','Mengunci versi PHP, Laravel, PhpSpreadsheet, dan paket pendukung.'),
('Frontend','npm; package-lock.json','Mengelola Vite, Tailwind CSS, dan tooling frontend.'),
('Pengujian','PHPUnit; tests/Unit; tests/Feature','Pengujian model, fitur, dan integrasi HTTP.'),
('Kualitas kode','Laravel Pint','Konsistensi format kode PHP.'),
('Build & rilis','npm run build; deploy.sh','Build aset production dan penyiapan deployment.')],[1800,3000,4560])
d.add_heading('6. Alur Keterkaitan Teknologi',1)
table(d,['Tahap','Alur','Komponen'],[
('1','Pengguna mengakses halaman dan mengirim formulir','Browser, Blade, CSS/Tailwind, route web, CSRF.'),
('2','Permintaan diperiksa dan diarahkan','Session, auth, role middleware, controller.'),
('3','Use case memproses data','Controller/service, validasi, Eloquent.'),
('4','Data/berkas disimpan','Database relasional, private storage, tabel domain SDM.'),
('5','Hasil disajikan atau diekspor','Blade, JSON internal, XLSX, ZIP, download response.'),
('6','Aktivitas dicatat dan diuji','activity_logs, PHPUnit, logging, pemeriksaan operasional.')],[1200,3900,4260])
d.add_heading('7. Catatan Teknologi Saat Ini',1)
label(d,'API publik: ','Belum tersedia. Proyek belum memiliki routes/api.php atau autentikasi Bearer token; endpoint saat ini adalah endpoint web internal berbasis session Laravel.')
label(d,'Database engine: ','Engine konkret mengikuti konfigurasi environment Laravel yang digunakan saat deployment.')
label(d,'Integrasi eksternal: ','Tidak ada integrasi layanan eksternal yang tercatat dalam ruang lingkup dokumentasi ini.')
label(d,'Sumber acuan: ','docs/arsitektur-teknologi.md, docs/arsitektur-aplikasi.md, docs/arsitektur-data.md, docs/security.md, docs/api.md, composer.json, package.json, dan docs/sistem-sdm-final.md.')
OUT.parent.mkdir(parents=True,exist_ok=True); d.save(OUT); print(OUT)

